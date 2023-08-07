import asyncio
import PyPDF2
from pprint import pprint
import re
from docx import Document
import functools
from pathlib import Path
from docx.oxml import OxmlElement


def _add_hyperlink():
    def decorator(func):
        @functools.wraps(func)
        async def wrapper(*args, **kwargs):
            doc, doc_name = await func(*args, **kwargs)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if re.search(r'https://', run.text):
                        hyperlink = run.text
                        run.text = ''
                        hyperlink_run = run._r
                        hyperlink_element = OxmlElement('w:hyperlink')
                        hyperlink_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id',
                                                run.part.relate_to(hyperlink_run, 'hyperlink'))
                        new_run = OxmlElement('w:r')
                        new_run.append(hyperlink_element)
                        new_run.text = hyperlink
                        print(hyperlink)
                        paragraph._p.insert(paragraph._p.index(hyperlink_run), new_run)
                        paragraph._p.remove(hyperlink_run)
            return doc, doc_name
        return wrapper
    return decorator





def _modify_properties():
    def decorator(func):
        @functools.wraps(func)
        async def wrapper(*args, **kwargs):
            doc, doc_name = await func(*args, **kwargs)
            
            location = doc_name[:doc_name.find('_')]
            doc_props = doc.core_properties
            properties = {
                    'Title': f'{location}',
                    'Author': 'yousefabuz17',
                    'Subject': 'Job URLs',
                    'Keywords': 'jobs, searching, grouping',
                    'Comments': 'Document is created based on the location provided, othewrise all will show',
                    'Last Modified By': 'yousefabuz17'
                }
            doc_props.title = properties.get('Title', '')
            doc_props.author = properties.get('Author', '')
            doc_props.subject = properties.get('Subject', '')
            doc_props.keywords = properties.get('Keywords', '')
            doc_props.comments = properties.get('Comments', '')
            doc_props.last_modified_by = properties.get('Last Modified By', '')
            return doc, doc_name
        return wrapper
    return decorator

def _export_to_docx():
    def decorator(func):
        @functools.wraps(func)
        async def wrapper(*args, **kwargs):
            pdf, location = await func(*args, **kwargs)
            doc = Document()
            
            if location is None:
                doc_name = 'job_urls.docx'
            else:
                doc_name = f'{location}_urls.docx'
            
            for idx, element in enumerate(pdf, start=1):
                doc.add_paragraph(f'{idx}: {element}\n')
            doc.save(doc_name)
            return doc, doc_name
        return wrapper
    return decorator

async def pdf_reader():
    with open(Path(__file__).parent.absolute() / 'job_urls.pdf', 'rb') as jobs:
        pdf = PyPDF2.PdfReader(jobs)
        all_text = [pdf.pages[i].extract_text().split('\n') for i in range(len(pdf.pages))]
    return all_text

# @_add_hyperlink()
@_modify_properties()
@_export_to_docx()
async def pdf_to_doc(pdf, location=None):
    if location is None:
        pdf = [j.strip() for i in pdf for j in i]
        return pdf, location
    filtered_loc = [j.strip() for i in pdf for j in i if re.search(rf'{location}', j)]
    return filtered_loc, location

async def main():
    pdf_text = await pdf_reader()
    docx = await pdf_to_doc(pdf_text, location=None)

if __name__ == '__main__':
    asyncio.run(main())